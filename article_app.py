import sys
if 'warnings' not in sys.modules:
    import warnings
    sys.modules['warnings'] = warnings

import streamlit as st
import os
import io
import re
import threading
import time
from datetime import datetime
from io import BytesIO
import pytz

# ─── SESSION STATE — MUST BE FIRST BEFORE CREWAI LOADS ───────────────────────
# CrewAI's internal EventsBus fires in a background thread immediately on import.
# If session state is not initialised before that happens, it throws:
# "st.session_state has no attribute raw_trace"
# Fix: initialise all session keys here, before any CrewAI import.
if "raw_trace" not in st.session_state:
    st.session_state.raw_trace = []
if "result" not in st.session_state:
    st.session_state.result = None
if "generating" not in st.session_state:
    st.session_state.generating = False

# ─── CONFIGURATION ────────────────────────────────────────────────────────────
os.environ["CREWAI_TRACING_ENABLED"] = "false"
os.environ["CREWAI_DISABLE_TELEMETRY"] = "true"

from crewai import Agent, Task, Crew, LLM

# ─── PAGE CONFIG ──────────────────────────────────────────────────────────────
st.set_page_config(page_title="Avernixx Article Intelligence Engine", layout="wide")

# ─── CUSTOM CSS ───────────────────────────────────────────────────────────────
st.markdown("""
<style>
    .main { max-width: 1100px; margin: auto; }
    .stTextArea textarea { white-space: pre-wrap !important; word-wrap: break-word !important; }

    .output-box {
        background: #0f1117;
        border: 1px solid #2e3a4e;
        border-radius: 10px;
        padding: 20px;
        white-space: pre-wrap;
        word-wrap: break-word;
        overflow-wrap: break-word;
        font-size: 0.95rem;
        line-height: 1.8;
        max-height: 680px;
        overflow-y: auto;
    }
    .terminal-box {
        background: #0a0f0a;
        border: 1px solid #1a3a1a;
        border-left: 4px solid #39ff14;
        border-radius: 8px;
        padding: 16px 20px;
        font-family: 'Courier New', Consolas, monospace;
        font-size: 0.80rem;
        color: #39ff14;
        white-space: pre-wrap;
        word-wrap: break-word;
        line-height: 1.55;
        max-height: 480px;
        overflow-y: auto;
    }
    .t-yellow { color: #FFD700; }
    .t-cyan   { color: #00FFFF; }
    .t-white  { color: #FFFFFF; }
    .t-red    { color: #FF6B6B; }
    .t-purple { color: #CF9FFF; }
    .trace-header {
        font-size: 0.75rem;
        color: #39ff14;
        letter-spacing: 2px;
        text-transform: uppercase;
        margin-bottom: 4px;
        font-family: 'Courier New', monospace;
    }
    .header-desc {
        background: linear-gradient(90deg, #0d1b2a, #1a2a3a);
        border-left: 4px solid #1E88E5;
        border-radius: 8px;
        padding: 14px 20px;
        margin: 10px 0 20px 0;
        font-size: 0.93rem;
        color: #cdd6e0;
        line-height: 1.7;
    }
    .llm-badge {
        display: inline-block;
        padding: 4px 14px;
        border-radius: 20px;
        font-size: 0.78rem;
        font-weight: 700;
        margin-bottom: 8px;
    }
    .badge-local  { background: #1b5e20; color: #69ff47; border: 1px solid #69ff47; }
    .badge-openai { background: #0a3d62; color: #74b9ff; border: 1px solid #74b9ff; }
    .disclaimer-bar {
        background: #1a1f2e;
        border: 1px solid #f4511e44;
        border-radius: 8px;
        padding: 10px 16px;
        text-align: center;
        font-size: 0.78rem;
        color: #aaa;
        margin-top: 8px;
    }
</style>
""", unsafe_allow_html=True)

# ─── STDOUT INTERCEPTOR ───────────────────────────────────────────────────────
class StreamCapture(io.StringIO):
    def __init__(self, real_stdout):
        super().__init__()
        self.real_stdout = real_stdout

    def write(self, text):
        self.real_stdout.write(text)
        self.real_stdout.flush()
        if text and text.strip():
            st.session_state.raw_trace.append(text)
        return len(text)

    def flush(self):
        self.real_stdout.flush()

# ─── ANSI STRIP & COLOUR ──────────────────────────────────────────────────────
ANSI_ESC = re.compile(r'\x1B(?:[@-Z\\-_]|\[[0-?]*[ -/]*[@-~])')

def to_html(lines):
    raw = "".join(lines)
    raw = ANSI_ESC.sub('', raw)
    raw = raw.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
    out = []
    for line in raw.split('\n'):
        if any(k in line for k in ['Task Completed','Final Answer','✅','COMPLETE']):
            out.append(f"<span class='t-cyan'>{line}</span>")
        elif any(k in line for k in ['Task Started','Agent Started','CREW','📋']):
            out.append(f"<span class='t-yellow'>{line}</span>")
        elif any(k in line for k in ['Thought:','Action:','Observation:']):
            out.append(f"<span class='t-purple'>{line}</span>")
        elif any(k in line for k in ['Error','error','Failed','❌']):
            out.append(f"<span class='t-red'>{line}</span>")
        elif any(k in line for k in ['Planner','Writer','Editor','Casey','Riley','Morgan']):
            out.append(f"<span class='t-white'>{line}</span>")
        else:
            out.append(line)
    return '\n'.join(out)

# ─── SINGAPORE TIME ───────────────────────────────────────────────────────────
def sg_timestamp():
    return datetime.now(pytz.timezone("Asia/Singapore")).strftime("%d %B %Y, %I:%M %p SGT")

def sg_stamp():
    return datetime.now(pytz.timezone("Asia/Singapore")).strftime("%Y%m%d_%I%M%p")

# ─── SESSION STATE already initialised at top of file ────────────────────────

# ─── HEADER ───────────────────────────────────────────────────────────────────
st.markdown("<h1 style='text-align:center;color:#1E88E5;letter-spacing:2px;'>⚡ AVERNIXX</h1>", unsafe_allow_html=True)
st.markdown("<h3 style='text-align:center;color:#cdd6e0;margin-top:-10px;'>Article Intelligence Engine — Multi-Agent Content System</h3>", unsafe_allow_html=True)

st.markdown("""
<div class='header-desc'>
<b>✍️ What This Agent Does:</b><br>
This is a <b>Multi-Agent Content System (MAS)</b> powered by <b>CrewAI</b> — part of Avernixx's
<b>Agentic AI Deployment Strategy</b>. Three specialised AI agents collaborate to research, write,
and polish a publication-ready article on any topic you define:<br><br>
&nbsp;&nbsp;🔵 <b>Strategist Casey</b> — Content Planner who researches trends, identifies your audience, builds the outline, and selects SEO keywords<br>
&nbsp;&nbsp;🟢 <b>Author Riley</b> — Content Writer who crafts a compelling, well-structured article based on Casey's plan<br>
&nbsp;&nbsp;🟣 <b>Editor Morgan</b> — Senior Editor who proofreads, balances viewpoints, aligns tone, and prepares the final publication-ready piece<br><br>
Powered by your choice of <b>local LLM (Ollama/Llama3)</b> or <b>OpenAI GPT</b> — no single prompt can match the depth of three collaborating agents.
</div>
""", unsafe_allow_html=True)

# ─── EXPANDERS ────────────────────────────────────────────────────────────────
with st.expander("🗺️ How the Multi-Agent Content System Works", expanded=False):
    st.markdown("""
```
┌─────────────────────────────────────────────────────────────┐
│            👤  USER INPUT (Topic + Controls)                │
└───────────────────────────┬─────────────────────────────────┘
                            ▼
┌─────────────────────────────────────────────────────────────┐
│           🔴  CREW ORCHESTRATOR (CrewAI Engine)             │
│    Receives inputs → Splits into tasks → Delegates          │
└──────────┬──────────────────┬──────────────────┬────────────┘
           ▼                  ▼                  ▼
   ┌──────────────┐   ┌──────────────┐   ┌──────────────┐
   │ 🔵 TASK 1    │   │ 🟢 TASK 2    │   │ 🟣 TASK 3    │
   │ Strategist   │   │ Author Riley │   │ Editor Morgan│
   │ Casey        │   │              │   │              │
   │ Research +   │   │ Write the    │   │ Proofread +  │
   │ Plan + SEO   │   │ full article │   │ Final polish │
   └──────┬───────┘   └──────┬───────┘   └──────┬───────┘
          └──────────────────┴──────────────────┘
                             ▼
┌─────────────────────────────────────────────────────────────┐
│    🧠  LLM BRAIN (Ollama/Llama3 or OpenAI GPT)              │
└───────────────────────────┬─────────────────────────────────┘
                            ▼
┌─────────────────────────────────────────────────────────────┐
│    📄  OUTPUT — Publication-Ready Article + Downloads       │
└─────────────────────────────────────────────────────────────┘
```
    """)

with st.expander("📘 LLM Options — Local vs Cloud", expanded=False):
    st.markdown("""
| | **Ollama / Llama3 (Local)** | **OpenAI GPT-3.5 / GPT-4** |
|---|---|---|
| **Cost** | Free — runs on your machine | Requires API key + usage fees |
| **Privacy** | 100% local — no data leaves your device | Data sent to OpenAI servers |
| **Speed** | Slower (depends on your hardware) | Faster |
| **Quality** | Good for demos | Higher quality output |
| **Internet** | Not required after setup | Required |
| **Best for** | Demos, privacy-sensitive use cases | Production, client-facing output |

> **Recommendation:** Use Ollama for demos. Switch to OpenAI for client deliverables.
    """)

# ─── LLM SELECTOR ─────────────────────────────────────────────────────────────
st.markdown("### 🧠 Select Your LLM")
llm_choice = st.radio(
    "Choose the AI brain for your agents:",
    ["🖥️  Local — Ollama / Llama3 (Free, Private)", "☁️  Cloud — OpenAI GPT (API Key Required)"],
    horizontal=True
)

openai_key = None
if "OpenAI" in llm_choice:
    st.markdown("<div class='llm-badge badge-openai'>☁️ OpenAI Mode</div>", unsafe_allow_html=True)
    openai_key = st.text_input("Enter your OpenAI API Key:", type="password", placeholder="sk-...")
    model_choice = st.selectbox("Select GPT Model:", ["gpt-3.5-turbo", "gpt-4", "gpt-4-turbo"])
else:
    st.markdown("<div class='llm-badge badge-local'>🖥️ Local Ollama Mode</div>", unsafe_allow_html=True)

# ─── ARTICLE CONTROLS ─────────────────────────────────────────────────────────
st.markdown("### 🎯 Article Configuration")

col1, col2 = st.columns(2)

with col1:
    topic = st.text_area(
        "📌 Article Topic",
        value="The impact of Agentic AI on enterprise digital transformation",
        height=80,
        help="Be as specific or broad as you like. The agents will research and frame it."
    )
    audience = st.selectbox(
        "👥 Target Audience",
        ["C-Suite / Executive Leadership", "IT & Technology Professionals",
         "Business Analysts & Managers", "General Public / Non-Technical",
         "Students & Academics", "Entrepreneurs & Startups"]
    )

with col2:
    tone = st.selectbox(
        "🎨 Writing Tone",
        ["Professional & Authoritative", "Conversational & Engaging",
         "Analytical & Data-Driven", "Inspirational & Thought Leadership",
         "Educational & Explanatory", "Persuasive & Opinion-Led"]
    )
    word_count = st.select_slider(
        "📏 Target Word Count",
        options=[300, 500, 800, 1000, 1200, 1500, 2000],
        value=800
    )
    seo_keywords = st.text_input(
        "🔍 SEO Keywords (optional)",
        placeholder="e.g. agentic AI, enterprise automation, LLM deployment",
        help="Comma-separated. Leave blank to let the agents decide."
    )

# ─── AGENTS ───────────────────────────────────────────────────────────────────
def build_llm():
    if "OpenAI" in llm_choice and openai_key:
        os.environ["OPENAI_API_KEY"] = openai_key
        return LLM(model=model_choice)
    else:
        return LLM(model="ollama/llama3", base_url="http://localhost:11434")

def build_agents(llm):
    casey = Agent(
        role="Strategist Casey — Content Planning & Research Lead",
        goal=f"Plan a compelling, SEO-optimised content strategy for an article on: {topic}",
        backstory=(
            "Senior Content Strategist at Avernixx with expertise in digital publishing, "
            "SEO architecture, and audience analysis. Casey has planned content strategies "
            "for Fortune 500 brands and enterprise AI publications. Known for turning complex "
            "topics into structured, audience-targeted content blueprints."
        ),
        allow_delegation=False,
        verbose=True,
        llm=llm
    )
    riley = Agent(
        role="Author Riley — Senior Content Writer",
        goal=f"Write a compelling, well-structured article on: {topic}",
        backstory=(
            "Principal Content Writer at Avernixx with a background in technology journalism "
            "and enterprise communications. Riley has written for leading AI and business "
            "publications, specialising in making technical topics accessible and engaging "
            "for diverse audiences. Known for sharp prose and strong narrative structure."
        ),
        allow_delegation=False,
        verbose=True,
        llm=llm
    )
    morgan = Agent(
        role="Editor Morgan — Senior Editorial & Publishing Lead",
        goal="Edit and finalise the article to publication-ready standard",
        backstory=(
            "Lead Editor at Avernixx with 12 years of editorial experience across "
            "enterprise technology, AI, and business media. Morgan ensures every piece "
            "meets journalistic best practices, maintains balanced viewpoints, aligns "
            "with brand voice, and is polished for immediate publication."
        ),
        allow_delegation=False,
        verbose=True,
        llm=llm
    )
    return casey, riley, morgan

# ─── BUTTONS ──────────────────────────────────────────────────────────────────
col_btn1, col_btn2 = st.columns([2, 1])
with col_btn1:
    generate_btn = st.button("✍️ Generate Article", use_container_width=True, type="primary")
with col_btn2:
    reset_btn = st.button("🔄 Reset / New Article", use_container_width=True)

if reset_btn:
    st.session_state.result     = None
    st.session_state.raw_trace  = []
    st.session_state.generating = False
    st.rerun()

# ─── TRACE & SPINNER SLOTS ────────────────────────────────────────────────────
trace_slot   = st.empty()
spinner_slot = st.empty()

def render_trace():
    if st.session_state.raw_trace:
        html  = to_html(st.session_state.raw_trace)
        label = "▶ LIVE AGENT TERMINAL — REAL-TIME CREW TRACE" if st.session_state.generating else "▶ AGENT TERMINAL LOG — SESSION COMPLETE"
        trace_slot.markdown(
            f"<div class='trace-header'>{label}</div>"
            f"<div class='terminal-box'>{html}</div>",
            unsafe_allow_html=True
        )

# ─── GENERATION ───────────────────────────────────────────────────────────────
if generate_btn and topic.strip():

    # Validate OpenAI key if needed
    if "OpenAI" in llm_choice and not openai_key:
        st.error("❌ Please enter your OpenAI API key above before generating.")
        st.stop()

    st.session_state.result     = None
    st.session_state.raw_trace  = []
    st.session_state.generating = True

    llm_label = f"OpenAI / {model_choice}" if "OpenAI" in llm_choice else "Ollama / Llama3 (Local)"

    banner = (
        "╔══════════════════════════════════════════════════════════════╗\n"
        "║    AVERNIXX ARTICLE INTELLIGENCE ENGINE — INITIALISING      ║\n"
        f"║    Topic:    {topic[:50]:<50}║\n"
        f"║    Tone:     {tone[:50]:<50}║\n"
        f"║    Audience: {audience[:50]:<50}║\n"
        f"║    Words:    {str(word_count):<50}║\n"
        f"║    LLM:      {llm_label[:50]:<50}║\n"
        "╚══════════════════════════════════════════════════════════════╝\n\n"
        "[CREW]  Assembling: Strategist Casey · Author Riley · Editor Morgan\n"
        "[CREW]  Mode: Sequential task execution\n"
        "──────────────────────────────────────────────────────────────\n\n"
    )
    st.session_state.raw_trace.append(banner)
    render_trace()

    real_stdout = sys.stdout
    sys.stdout  = StreamCapture(real_stdout)

    result_holder = {}
    error_holder  = {}

    seo_note = f"Prioritise these SEO keywords: {seo_keywords}." if seo_keywords.strip() else "Identify the most relevant SEO keywords for this topic."

    def run_crew():
        try:
            llm = build_llm()
            casey, riley, morgan = build_agents(llm)

            t1 = Task(
                description=(
                    f"You are Strategist Casey, Content Planning Lead at Avernixx.\n"
                    f"Research and plan a content strategy for an article on: '{topic}'.\n\n"
                    f"Your plan must include:\n"
                    f"1. Latest trends and key developments on this topic\n"
                    f"2. Target audience analysis — their interests, pain points, and knowledge level\n"
                    f"   (Target audience for this article: {audience})\n"
                    f"3. A detailed content outline: introduction, 4-6 key sections, conclusion, and call to action\n"
                    f"4. {seo_note}\n"
                    f"5. Recommended sources, data points, or references to strengthen credibility\n"
                    f"6. Tone guidance for the writer (Tone: {tone})"
                ),
                expected_output=(
                    "A comprehensive content plan with: audience analysis, full article outline, "
                    "SEO keywords, tone guidance, and recommended sources."
                ),
                agent=casey
            )

            t2 = Task(
                description=(
                    f"You are Author Riley, Senior Content Writer at Avernixx.\n"
                    f"Using Strategist Casey's content plan, write a complete article on: '{topic}'.\n\n"
                    f"Requirements:\n"
                    f"1. Target word count: approximately {word_count} words\n"
                    f"2. Tone: {tone}\n"
                    f"3. Audience: {audience}\n"
                    f"4. Incorporate SEO keywords naturally — never forced\n"
                    f"5. Structure: engaging introduction, well-named sections with subheadings, "
                    f"   insightful body paragraphs (2-3 per section), and a strong conclusion\n"
                    f"6. Each section should have 2-3 paragraphs\n"
                    f"7. Clearly distinguish between facts and opinions\n"
                    f"8. Write in Markdown format"
                ),
                expected_output=(
                    f"A complete, well-structured article of approximately {word_count} words "
                    f"in Markdown format, ready for editorial review."
                ),
                agent=riley
            )

            t3 = Task(
                description=(
                    f"You are Editor Morgan, Senior Editorial Lead at Avernixx.\n"
                    f"Review and finalise Author Riley's article on: '{topic}'.\n\n"
                    f"Your editorial tasks:\n"
                    f"1. Proofread for grammatical errors, typos, and inconsistencies\n"
                    f"2. Ensure tone consistency throughout ({tone})\n"
                    f"3. Verify the article is appropriate for the target audience: {audience}\n"
                    f"4. Ensure balanced viewpoints — flag and soften any one-sided assertions\n"
                    f"5. Confirm SEO keywords are naturally integrated\n"
                    f"6. Sharpen the introduction and conclusion if needed\n"
                    f"7. Ensure the article is publication-ready\n"
                    f"8. Return the final article in clean Markdown format"
                ),
                expected_output=(
                    "A polished, publication-ready article in Markdown format "
                    "with all editorial improvements applied."
                ),
                agent=morgan
            )

            crew = Crew(
                agents=[casey, riley, morgan],
                tasks=[t1, t2, t3],
                verbose=True
            )
            result_holder["output"] = str(crew.kickoff())
        except Exception as e:
            error_holder["error"] = str(e)

    thread = threading.Thread(target=run_crew, daemon=True)
    thread.start()

    while thread.is_alive():
        render_trace()
        spinner_slot.info("✍️ Agents are writing your article — watch the live trace above. Takes 1–3 min on local LLM.")
        time.sleep(1.5)
        st.rerun()

    sys.stdout = real_stdout

    st.session_state.raw_trace.append(
        "\n──────────────────────────────────────────────────────────────\n"
        f"[CREW]  ✅  ARTICLE COMPLETE — Ready for publication.\n"
        f"[CREW]  Timestamp: {sg_timestamp()}\n"
    )
    st.session_state.generating = False
    spinner_slot.empty()

    if "error" in error_holder:
        st.error(f"❌ Crew error: {error_holder['error']}")
    else:
        st.session_state.result = result_holder.get("output", "")

    render_trace()

# ─── PERSISTENT TRACE ─────────────────────────────────────────────────────────
if st.session_state.raw_trace and not st.session_state.generating and not generate_btn:
    render_trace()

# ─── OUTPUT ───────────────────────────────────────────────────────────────────
if st.session_state.result:
    st.markdown("---")
    st.markdown("### 📄 Your Generated Article")
    safe = st.session_state.result.replace("<","&lt;").replace(">","&gt;")
    st.markdown(f"<div class='output-box'>{safe}</div>", unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("### 📥 Download Your Article")

    stamp    = sg_stamp()
    ts_label = sg_timestamp()
    topic_slug = topic[:40].replace(" ","_").replace("/","_")
    header   = (
        f"# Article: {topic}\n"
        f"Generated by: Avernixx Article Intelligence Engine\n"
        f"Timestamp: {ts_label}\n"
        f"Tone: {tone}  |  Audience: {audience}  |  Words: ~{word_count}\n\n---\n\n"
    )
    full_md = header + st.session_state.result

    col_a, col_b, col_c = st.columns(3)

    with col_a:
        st.download_button(
            "📄 Download Markdown (.md)", data=full_md,
            file_name=f"Avernixx_Article_{topic_slug}_{stamp}.md",
            mime="text/markdown", use_container_width=True
        )
    with col_b:
        st.download_button(
            "📝 Download Text (.txt)", data=full_md,
            file_name=f"Avernixx_Article_{topic_slug}_{stamp}.txt",
            mime="text/plain", use_container_width=True
        )
    with col_c:
        try:
            from docx import Document as DocxDoc
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = DocxDoc()
            h = doc.add_heading("Avernixx Article Intelligence Engine", 0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            meta_lines = [
                f"Topic: {topic}",
                f"Generated: {ts_label}",
                f"Tone: {tone}  |  Audience: {audience}  |  Target Words: ~{word_count}"
            ]
            for line in meta_lines:
                mp = doc.add_paragraph(line)
                mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                mp.runs[0].font.size = Pt(10)
                mp.runs[0].font.color.rgb = RGBColor(0x66,0x66,0x66)
            doc.add_paragraph()

            for line in st.session_state.result.split("\n"):
                line = line.strip()
                if   line.startswith("### "): doc.add_heading(line[4:], level=3)
                elif line.startswith("## "):  doc.add_heading(line[3:], level=2)
                elif line.startswith("# "):   doc.add_heading(line[2:], level=1)
                elif line.startswith(("- ","* ")): doc.add_paragraph(line[2:], style="List Bullet")
                elif line == "---": doc.add_paragraph("─" * 60)
                elif line: doc.add_paragraph(line)

            doc.add_paragraph()
            fp = doc.add_paragraph(
                "⚠️ DISCLAIMER: Generated by Avernixx Article Intelligence Engine AI Demo. "
                "For demonstration purposes only. Validate all content before publication."
            )
            fp.runs[0].font.size = Pt(8)
            fp.runs[0].font.color.rgb = RGBColor(0x99,0x99,0x99)

            buf = BytesIO(); doc.save(buf); buf.seek(0)
            st.download_button(
                "📘 Download Word (.docx)", data=buf,
                file_name=f"Avernixx_Article_{topic_slug}_{stamp}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        except ImportError:
            st.warning("Run: pip install python-docx")

    st.markdown(f"<div style='text-align:center;color:#666;font-size:0.8rem;margin-top:6px;'>All files timestamped: {ts_label}</div>", unsafe_allow_html=True)

# ─── FOOTER ───────────────────────────────────────────────────────────────────
st.markdown("<br>", unsafe_allow_html=True)
st.markdown("""
<hr>
<div style='text-align:center;color:gray;font-size:0.85rem;'>
    © 2026 Avernixx Pte. Ltd. | <a href='http://www.avernixx.com' style='color:#1E88E5;'>www.avernixx.com</a>
</div>
<div class='disclaimer-bar'>
    ⚠️ <b>DEMO DISCLAIMER:</b> This system is a technology demonstration of Agentic AI content capabilities.
    All outputs are AI-generated and intended for exploratory and educational purposes only.
    Always review and validate AI-generated content before publication.
    Avernixx Pte. Ltd. accepts no liability for published content based on this demo output.
</div>
""", unsafe_allow_html=True)
